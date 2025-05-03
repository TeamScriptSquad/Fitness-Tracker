import os
from datetime import datetime
from aiogram import Bot, types
from aiogram.types import Message, FSInputFile
from aiogram.fsm.context import FSMContext
import pandas as pd
import logging

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Проверка доступности библиотек
try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("Создание Word-документов недоступно (установите python-docx)")

try:
    import openpyxl

    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logger.warning("Создание Excel-файлов недоступно (установите openpyxl)")


class TemplateStates:
    WAITING_FOR_TEMPLATE_NAME = "waiting_for_template_name"
    WAITING_FOR_EXERCISES = "waiting_for_exercises"
    WAITING_FOR_FORMAT = "waiting_for_format"


async def start_template_creation(message: Message, state: FSMContext):
    """Начинает процесс создания шаблона"""
    await message.answer("📝 Введите название для нового шаблона тренировки:")
    await state.set_state(TemplateStates.WAITING_FOR_TEMPLATE_NAME)


async def process_template_name(message: Message, state: FSMContext):
    """Обрабатывает название шаблона"""
    await state.update_data(template_name=message.text)
    await message.answer(
        "🏋️ Теперь введите упражнения в формате:\n"
        "<b>Название упражнения, подходы, повторения, вес</b>\n"
        "Каждое упражнение с новой строки\n\n"
        "Пример:\n"
        "Жим лежа, 4, 10, 60\n"
        "Приседания, 3, 12, 50",
        parse_mode="HTML"
    )
    await state.set_state(TemplateStates.WAITING_FOR_EXERCISES)


async def process_exercises_data(message: Message, state: FSMContext, bot: Bot):
    """Обрабатывает введенные упражнения"""
    data = await state.get_data()
    template_name = data.get('template_name')
    exercises_text = message.text

    try:
        exercises = []
        for line in exercises_text.split('\n'):
            line = line.strip()
            if line:
                parts = [p.strip() for p in line.split(',')]
                if len(parts) != 4:
                    raise ValueError("❌ Неверный формат строки. Нужно 4 значения через запятую")

                name, sets, reps, weight = parts
                exercises.append({
                    'Упражнение': name,
                    'Подходы': int(sets),
                    'Повторения': int(reps),
                    'Вес (кг)': float(weight)
                })

        await state.update_data(exercises=exercises)

        # Формируем клавиатуру
        buttons = []
        if EXCEL_AVAILABLE:
            buttons.append([types.KeyboardButton(text="Excel")])
        if DOCX_AVAILABLE:
            buttons.append([types.KeyboardButton(text="Word")])
        if EXCEL_AVAILABLE and DOCX_AVAILABLE:
            buttons.append([types.KeyboardButton(text="Оба формата")])

        if not buttons:
            await message.answer("❌ Нет доступных форматов (установите openpyxl и/или python-docx)")
            await state.clear()
            return

        await message.answer(
            "📁 Выберите формат для сохранения:",
            reply_markup=types.ReplyKeyboardMarkup(
                keyboard=buttons,
                resize_keyboard=True
            )
        )
        await state.set_state(TemplateStates.WAITING_FOR_FORMAT)

    except Exception as e:
        await message.answer(f"❌ Ошибка: {str(e)}")
        logger.error(f"Ошибка обработки упражнений: {str(e)}")
        await state.clear()


async def process_format_selection(message: Message, state: FSMContext, bot: Bot):
    """Обрабатывает выбор формата и создает файлы"""
    data = await state.get_data()
    template_name = data.get('template_name')
    exercises = data.get('exercises')
    format_choice = message.text.lower()

    try:
        files_created = []

        # Создание Excel файла
        if format_choice in ["excel", "оба формата"] and EXCEL_AVAILABLE:
            excel_file = f"{template_name}.xlsx"
            try:
                df = pd.DataFrame(exercises)
                df.to_excel(excel_file, index=False, engine='openpyxl')
                files_created.append(excel_file)
                logger.info(f"Создан Excel-файл: {excel_file}")
            except Exception as e:
                logger.error(f"Ошибка создания Excel: {str(e)}")
                await message.answer("❌ Не удалось создать Excel-файл")

        # Создание Word файла
        if format_choice in ["word", "оба формата"] and DOCX_AVAILABLE:
            word_file = f"{template_name}.docx"
            try:
                doc = Document()
                doc.add_heading(f"Шаблон: {template_name}", 0)

                # Таблица
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text = 'Упражнение'
                hdr[1].text = 'Подходы'
                hdr[2].text = 'Повторения'
                hdr[3].text = 'Вес (кг)'

                for ex in exercises:
                    row = table.add_row().cells
                    row[0].text = ex['Упражнение']
                    row[1].text = str(ex['Подходы'])
                    row[2].text = str(ex['Повторения'])
                    row[3].text = str(ex['Вес (кг)'])

                doc.save(word_file)
                files_created.append(word_file)
                logger.info(f"Создан Word-файл: {word_file}")
            except Exception as e:
                logger.error(f"Ошибка создания Word: {str(e)}")
                await message.answer("❌ Не удалось создать Word-документ")

        # Отправка и удаление файлов
        for file_path in files_created:
            try:
                # Отправка файла
                await bot.send_document(
                    chat_id=message.chat.id,
                    document=FSInputFile(file_path),
                    caption=f"🏋️ Шаблон '{template_name}'"
                )
                logger.info(f"Файл {file_path} отправлен")

            except Exception as e:
                logger.error(f"Ошибка отправки: {str(e)}")
                await message.answer(f"⚠️ Ошибка при отправке файла {file_path}")

            finally:
                # Гарантированное удаление файла
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                        logger.info(f"Файл {file_path} удален")
                    else:
                        logger.warning(f"Файл {file_path} не найден для удаления")
                except Exception as e:
                    logger.error(f"Ошибка удаления файла {file_path}: {str(e)}")

        if files_created:
            await message.answer("✅ Файлы успешно созданы и отправлены!")
        else:
            await message.answer("❌ Не удалось создать ни один файл")

        await state.clear()

    except Exception as e:
        logger.error(f"Критическая ошибка: {str(e)}")
        await message.answer("⚠️ Произошла критическая ошибка при обработке запроса")
        await state.clear()


async def send_template(message: Message):
    """Обработчик для команды /get_template"""
    await message.answer("ℹ️ Для получения шаблона используйте /create_template")